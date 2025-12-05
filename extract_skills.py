import os
import re
import tempfile
import logging
import asyncio
import concurrent.futures
import time
from functools import lru_cache
from fastapi import APIRouter, UploadFile, File
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from typing import List
from io import BytesIO
from PIL import Image
import spacy
import PyPDF2
from PyPDF2 import PdfReader
import gc

# ---------- CONFIG ----------
pytesseract.pytesseract.tesseract_cmd = os.getenv("TESSERACT_CMD", "/usr/bin/tesseract")
POPPLER_PATH = os.getenv("POPPLER_PATH", "/usr/bin")

# ---------- OPTIMIZATION CONFIG ----------
PROCESSING_CONFIG = {
    "max_workers": min(4, os.cpu_count() or 1),
    "max_pdf_pages": 50,
    "chunk_size": 10,
    "timeout": 300,
}

# ---------- LOGGING CONFIG ----------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s",
)
logger = logging.getLogger("ocr_nlp_logger")

router = APIRouter()

# ---------- SKILLS / FRAMEWORKS ----------
SKILL_KEYWORDS = [
    "Python", "Java", "JavaScript", "React", "SQL", "NoSQL",
    "API Design", "Project Management", "Agile", "Scrum", "System Architecture",
    "Leadership", "Teamwork", "Communication", "Machine Learning"
]

SKILL_SYNONYMS = {
    "JS": "JavaScript",
    "PM": "Project Management",
    "ML": "Machine Learning",
    "DB": "SQL",
}

COMMON_FRAMEWORKS = [
    "Django", "Flask", "FastAPI", "Vue", "Angular",
    "Node.js", "Express", "Spring Boot"
]

TECH_TERMS = {
    "Python", "Java", "JavaScript", "React", "SQL", "NoSQL",
    "Django", "Flask", "FastAPI", "Vue", "Angular", "Node.js", "Express",
    "Spring Boot", "PyTorch", "TensorFlow", "Keras", "Docker", "Kubernetes",
    "Terraform", "AWS", "GCP", "Azure", "MongoDB", "PostgreSQL", "MySQL"
}

# Combine all skill-related terms for better matching
ALL_SKILLS_SET = set(SKILL_KEYWORDS) | set(COMMON_FRAMEWORKS) | TECH_TERMS | set(SKILL_SYNONYMS.values())

# ---------- PRE-COMPILED REGEX PATTERNS ----------
HEADING_PATTERNS = [re.compile(rf"\b{re.escape(h)}\b", re.IGNORECASE) for h in [
    "Personal Information", "Education", "Skills", "Experience",
    "Work Experience", "Projects", "Certifications", "Summary"
]]

MULTI_NEWLINE = re.compile(r'\n+')
MULTI_SPACES = re.compile(r'\s{2,}')
EMAIL_PATTERN = re.compile(r"[\w\.-]+@[\w\.-]+")
PHONE_PATTERN = re.compile(r"(\+?\d[\d\s\-]{8,}\d)")
FULL_NAME_PATTERN = re.compile(r"Full Name[:\s]*(.+)", re.IGNORECASE)
EMPLOYEE_ID_PATTERN = re.compile(r"(Employee ID|ID)[:\s]*(.+)", re.IGNORECASE)
LOCATION_PATTERN = re.compile(r"Location[:\s]*(.+)", re.IGNORECASE)

# Improved skill patterns with case-insensitive matching
SKILL_PATTERNS = {}
ALL_SKILL_TERMS = SKILL_KEYWORDS + COMMON_FRAMEWORKS + list(SKILL_SYNONYMS.keys())
for term in ALL_SKILL_TERMS:
    SKILL_PATTERNS[term] = re.compile(r'\b' + re.escape(term) + r'\b', re.IGNORECASE)

# ---------- CACHED NLP MODEL ----------
@lru_cache(maxsize=1)
def get_nlp_model():
    """Cache the NLP model to avoid reloading"""
    logger.info("Loading spaCy model...")
    return spacy.load("en_core_web_sm")

nlp = get_nlp_model()

# ------------------------------------------------------
#   TIMING DECORATOR FOR DEBUGGING
# ------------------------------------------------------
def timing_decorator(func_name=""):
    def decorator(func):
        def wrapper(*args, **kwargs):
            start_time = time.time()
            logger.info(f"⏱️  STARTING {func_name or func.__name__}...")
            
            result = func(*args, **kwargs)
            
            end_time = time.time()
            duration = end_time - start_time
            logger.info(f"⏱️  COMPLETED {func_name or func.__name__} in {duration:.2f} seconds")
            
            return result
        return wrapper
    return decorator

# ------------------------------------------------------
#   FIXED DUAL APPROACH: PDF TEXT EXTRACTION WITH PROPER FILE HANDLING
# ------------------------------------------------------
@timing_decorator("PDF Text Extraction")
def extract_text_from_pdf_fixed(pdf_path):
    """Extract text from PDF using both direct extraction and OCR fallback with proper file handling"""
    logger.info(f"Starting dual PDF extraction for: {pdf_path}")
    
    text = ""
    
    # First attempt: Direct text extraction (for text-based PDFs)
    direct_text = ""
    try:
        logger.info("Attempting direct text extraction from PDF...")
        with open(pdf_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            
            for i, page in enumerate(pdf_reader.pages):
                if i >= PROCESSING_CONFIG["max_pdf_pages"]:
                    break
                page_text = page.extract_text()
                if page_text.strip():
                    direct_text += page_text + "\n"
            
            logger.info(f"Direct extraction got {len(direct_text)} characters")
            
    except Exception as e:
        logger.info(f"Direct extraction failed: {e}")
        direct_text = ""

    # Check if direct extraction got meaningful text
    if len(direct_text.strip()) > 100:  # If we got substantial text
        text = direct_text
        logger.info("Using direct text extraction (text-based PDF)")
    else:
        logger.info("Direct extraction insufficient - trying OCR for scanned PDF...")
        
        # Second attempt: OCR extraction (for scanned PDFs)
        try:
            images = convert_from_path(
                pdf_path, 
                poppler_path=POPPLER_PATH,
                first_page=1, 
                last_page=PROCESSING_CONFIG["max_pdf_pages"],
                dpi=300,
                grayscale=True
            )
            logger.info(f"PDF converted into {len(images)} image pages for OCR")

            # Process pages with OCR
            ocr_text = ""
            for i, img in enumerate(images):
                logger.debug(f"OCR processing page {i + 1}")
                
                # Use optimized OCR configuration
                page_text = pytesseract.image_to_string(
                    img, 
                    config='--psm 6 -c preserve_interword_spaces=1',
                    lang='eng'
                )
                
                ocr_text += page_text + "\n"
                logger.debug(f"OCR page {i+1} extracted {len(page_text)} characters")
                
                # Explicitly clean up image to free memory
                del img
                if i % 5 == 0:  # Force garbage collection periodically
                    gc.collect()
            
            text = ocr_text
            logger.info(f"OCR extraction completed with {len(text)} characters")
            
        except Exception as ocr_error:
            logger.error(f"OCR extraction also failed: {ocr_error}")
            text = ""

    # Final check and debug info
    if text.strip():
        logger.info(f"PDF extraction successful. Total characters: {len(text)}")
        # Log sample for debugging
        sample = text[:500].replace('\n', ' ').strip()
        logger.debug(f"Extracted text sample: {sample}...")
        
        # Check if skills are likely to be found
        skill_keywords_found = [skill for skill in ALL_SKILLS_SET if skill.lower() in text.lower()]
        logger.info(f"Potential skills detected in text: {len(skill_keywords_found)}")
        if skill_keywords_found:
            logger.debug(f"Sample detected skills: {skill_keywords_found[:5]}")
    else:
        logger.warning("No text extracted from PDF using either method!")
        
    return text.strip()

# ------------------------------------------------------
#   OPTIMIZED DOCX TEXT EXTRACTION
# ------------------------------------------------------
@timing_decorator("DOCX Text Extraction")
def extract_text_from_docx_optimized(docx_file):
    logger.info("Starting optimized DOCX text extraction")
    text_parts = []

    try:
        doc = Document(docx_file)

        # Process paragraphs
        text_parts.extend(p.text for p in doc.paragraphs if p.text.strip())

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        text = "\n".join(text_parts)

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        text = ""

    logger.info(f"Finished DOCX extraction. Characters: {len(text)}")
    return text

# ------------------------------------------------------
#   IMPROVED TEXT CLEANING
# ------------------------------------------------------
@timing_decorator("Text Cleaning")
def clean_ocr_text_improved(text):
    logger.debug("Cleaning extracted text (improved)")
    
    # Remove headings using pre-compiled patterns
    for pattern in HEADING_PATTERNS:
        text = pattern.sub("", text)
    
    # Clean whitespace but preserve structure better
    text = MULTI_NEWLINE.sub('\n', text)
    text = MULTI_SPACES.sub(' ', text)
    
    return text.strip()

# ------------------------------------------------------
#   IMPROVED PERSONAL INFO EXTRACTION
# ------------------------------------------------------
@timing_decorator("Personal Info Extraction")
def extract_personal_info_improved(text):
    logger.info("Extracting personal info (improved)")
    info = {}
    
    lines = text.split("\n")

    for line in lines:
        line_lower = line.lower()

        if "full name" in line_lower:
            match = FULL_NAME_PATTERN.search(line)
            if match:
                info["Full Name"] = match.group(1).strip()

        elif "employee id" in line_lower or " id" in line_lower:
            match = EMPLOYEE_ID_PATTERN.search(line)
            if match:
                info["Employee ID"] = match.group(2).strip()

        elif "email" in line_lower:
            match = EMAIL_PATTERN.search(line)
            if match:
                info["Email"] = match.group()

        elif "phone" in line_lower:
            match = PHONE_PATTERN.search(line)
            if match:
                info["Phone Number"] = match.group()

        elif "location" in line_lower:
            match = LOCATION_PATTERN.search(line)
            if match:
                info["Location"] = match.group(1).strip()

    logger.debug("Applying NLP fallback for personal info")
    # Use original text for better NLP results
    doc = nlp(text[:100000])

    # Extract entities in single pass
    entities = {}
    for ent in doc.ents:
        if ent.label_ not in entities:
            entities[ent.label_] = ent.text
    
    if "Full Name" not in info and "PERSON" in entities:
        info["Full Name"] = entities["PERSON"]

    if "Location" not in info and "GPE" in entities:
        info["Location"] = entities["GPE"]

    if "Organization" not in info and "ORG" in entities:
        info["Organization"] = entities["ORG"]

    logger.info(f"Extracted personal info: {info}")
    return info

# ------------------------------------------------------
#   ROBUST SKILL EXTRACTION
# ------------------------------------------------------
@timing_decorator("Skill Extraction")
def extract_skills_robust(text):
    logger.info("Starting robust skill extraction")
    found_skills = set()
    
    if not text.strip():
        logger.warning("No text provided for skill extraction")
        return []

    # Debug: Check if common terms exist in text
    common_terms_in_text = [term for term in ['Python', 'Java', 'SQL', 'React'] 
                           if term.lower() in text.lower()]
    logger.debug(f"Common terms found in text: {common_terms_in_text}")

    # Method 1: Direct case-insensitive search (fastest)
    text_lower = text.lower()
    for skill in ALL_SKILLS_SET:
        if skill.lower() in text_lower:
            # More precise check with word boundaries
            pattern = re.compile(r'\b' + re.escape(skill.lower()) + r'\b', re.IGNORECASE)
            if pattern.search(text):
                found_skills.add(skill)
                logger.debug(f"Skill found (direct): {skill}")

    # Method 2: Regex pattern matching
    for term, pattern in SKILL_PATTERNS.items():
        if pattern.search(text) and SKILL_SYNONYMS.get(term, term) not in found_skills:
            actual_term = SKILL_SYNONYMS.get(term, term)
            found_skills.add(actual_term)
            logger.debug(f"Skill found (regex): {actual_term}")

    # Method 3: Check for multi-word skills
    multiword_skills = [skill for skill in ALL_SKILLS_SET if ' ' in skill]
    for skill in multiword_skills:
        if skill.lower() in text_lower and skill not in found_skills:
            found_skills.add(skill)
            logger.debug(f"Skill found (multi-word): {skill}")

    # Method 4: NLP-based extraction as final fallback
    if len(found_skills) < 3:  # If we found very few skills, try NLP
        logger.debug("Trying NLP-based skill extraction as fallback")
        nlp_text = text if len(text) < 30000 else text[:30000]
        doc = nlp(nlp_text)
        
        for token in doc:
            if token.text in ALL_SKILLS_SET and token.text not in found_skills:
                found_skills.add(token.text)
                logger.debug(f"Skill found (NLP): {token.text}")

    result = sorted(found_skills)
    logger.info(f"Final extracted skills: {result} (total: {len(result)})")
    return result

# ------------------------------------------------------
#   FIXED FILE PROCESSING WORKER WITH PROPER FILE CLEANUP AND TIMING
# ------------------------------------------------------
async def process_single_file_fixed(file: UploadFile):
    """Process a single file asynchronously with proper file handling"""
    temp_file_path = None
    file_start_time = time.time()
    
    try:
        logger.info(f"📁 STARTING FILE PROCESSING: {file.filename}")
        content = await file.read()
        suffix = os.path.splitext(file.filename)[1].lower()
        text = ""

        if suffix == ".pdf":
            logger.info(f"Handling PDF file: {file.filename}")
            
            # Create temporary file with explicit cleanup
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                tmp_pdf.write(content)
                tmp_pdf.flush()
                temp_file_path = tmp_pdf.name
            
            # Extract text from the temporary file
            text = extract_text_from_pdf_fixed(temp_file_path)
            
            # Explicitly close and delete the temporary file
            if temp_file_path and os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
                temp_file_path = None

        elif suffix == ".docx":
            logger.info(f"Handling DOCX file: {file.filename}")
            docx_file = BytesIO(content)
            text = extract_text_from_docx_optimized(docx_file)

        elif suffix in [".png", ".jpg", ".jpeg"]:
            logger.info(f"Handling image file: {file.filename}")

            # IMPORTANT: Convert to RGB always
            img = Image.open(BytesIO(content)).convert("RGB")

            # OCR in thread to avoid blocking
            with concurrent.futures.ThreadPoolExecutor() as executor:
                text = await asyncio.get_event_loop().run_in_executor(
                    executor, 
                    lambda: pytesseract.image_to_string(
                        img,
                        lang="eng",
                        config="--psm 6 -c preserve_interword_spaces=1"
                    )
                )

            logger.debug(f"Image OCR text length: {len(text)}")

        else:
            logger.warning(f"Unsupported file type: {suffix}")
            return {
                "filename": file.filename,
                "personal_info": {},
                "skills": []
            }

        if not text.strip():
            logger.warning(f"No text extracted from file: {file.filename}")
            return {
                "filename": file.filename,
                "personal_info": {},
                "skills": []
            }

        # Debug: Log extracted text characteristics
        logger.info(f"Extracted {len(text)} characters from {file.filename}")
        
        # Process personal info and skills
        personal_info = extract_personal_info_improved(text)
        skills = extract_skills_robust(text)
        
        file_end_time = time.time()
        file_duration = file_end_time - file_start_time
        logger.info(f"✅ COMPLETED FILE: {file.filename} in {file_duration:.2f} seconds")

        return {
            "filename": file.filename,
            "personal_info": personal_info,
            "skills": skills,
            "processing_time_seconds": round(file_duration, 2)
        }

    except Exception as e:
        file_end_time = time.time()
        file_duration = file_end_time - file_start_time
        logger.error(f"❌ ERROR processing file {file.filename} after {file_duration:.2f} seconds: {e}", exc_info=True)
        # Ensure temp file is cleaned up even if there's an error
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except Exception as cleanup_error:
                logger.warning(f"Could not cleanup temp file {temp_file_path}: {cleanup_error}")
        return {
            "filename": file.filename,
            "personal_info": {},
            "skills": [],
            "processing_time_seconds": round(file_duration, 2),
            "error": str(e)
        }

# ------------------------------------------------------
#   FIXED API ROUTE WITH COMPREHENSIVE TIMING
# ------------------------------------------------------
@router.post("/extract_skills/")
async def extract_skills_endpoint_fixed(files: List[UploadFile] = File(...)):
    total_start_time = time.time()
    logger.info(f"🚀 API /extract_skills called with {len(files)} files")
    
    # Process files concurrently
    tasks = [process_single_file_fixed(file) for file in files]
    results = await asyncio.gather(*tasks)
    
    # Calculate timing statistics
    total_end_time = time.time()
    total_duration = total_end_time - total_start_time
    
    # Extract all unique skills and timing info
    all_skills = sorted(set(
        skill for r in results for skill in r["skills"]
    ))
    
    # Calculate processing statistics
    successful_files = [r for r in results if r.get("skills")]
    failed_files = [r for r in results if not r.get("skills")]
    total_processing_time = sum(r.get("processing_time_seconds", 0) for r in results)
    avg_processing_time = total_processing_time / len(results) if results else 0

    logger.info("=" * 60)
    logger.info("📊 EXTRACTION SUMMARY:")
    logger.info(f"   Total files processed: {len(results)}")
    logger.info(f"   Successful extractions: {len(successful_files)}")
    logger.info(f"   Failed extractions: {len(failed_files)}")
    logger.info(f"   Total unique skills found: {len(all_skills)}")
    logger.info(f"   Total API processing time: {total_duration:.2f} seconds")
    logger.info(f"   Average file processing time: {avg_processing_time:.2f} seconds")
    logger.info(f"   Total file processing time: {total_processing_time:.2f} seconds")
    
    # Log individual file timings
    logger.info("   --- Individual File Timings ---")
    for result in results:
        filename = result.get("filename", "Unknown")
        processing_time = result.get("processing_time_seconds", 0)
        skills_count = len(result.get("skills", []))
        status = "✅ SUCCESS" if skills_count > 0 else "❌ FAILED"
        logger.info(f"   📄 {filename}: {status} - {skills_count} skills - {processing_time:.2f}s")
    logger.info("=" * 60)

    # Add timing information to response
    response = {
        "results": results, 
        "skills": all_skills,
        "processing_stats": {
            "total_files": len(results),
            "successful_files": len(successful_files),
            "failed_files": len(failed_files),
            "total_unique_skills": len(all_skills),
            "total_processing_time_seconds": round(total_duration, 2),
            "average_file_processing_time_seconds": round(avg_processing_time, 2),
            "individual_file_times": {
                r["filename"]: r.get("processing_time_seconds", 0) for r in results
            }
        }
    }
    
    logger.info(f"🎯 RETURNING RESPONSE after {total_duration:.2f} seconds")
    return response